"""
Document processing service.

Handles extension-specific text extraction (PDF, DOCX, TXT, CSV) and
orchestrates the file -> text -> cleaned text pipeline. Chunking, embedding,
and vector storage are delegated to their respective services from
rag_service, which composes all of these.
"""

import csv
import io
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from fastapi import HTTPException, status

from app.core.logging import logger
from app.utils.file_utils import clean_text


class DocumentService:
    """Extracts clean plain text from supported document formats."""

    def extract_text(self, file_path: Path, extension: str) -> str:
        extractors = {
            ".pdf": self._extract_pdf,
            ".docx": self._extract_docx,
            ".txt": self._extract_txt,
            ".csv": self._extract_csv,
        }

        extractor = extractors.get(extension)
        if extractor is None:
            raise HTTPException(
                status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
                detail=f"No extractor available for extension '{extension}'",
            )

        try:
            raw_text = extractor(file_path)
        except Exception as exc:  # noqa: BLE001 - we want to convert any parser error to HTTP 422
            logger.error(f"Failed to extract text from {file_path}: {exc}")
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail=f"Failed to extract text from file: {exc}",
            ) from exc

        text = clean_text(raw_text)
        if not text:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="No extractable text found in document",
            )
        return text

    @staticmethod
    def _extract_pdf(file_path: Path) -> str:
        text_parts: list[str] = []
        with fitz.open(file_path) as pdf:
            for page_number, page in enumerate(pdf, start=1):
                page_text = page.get_text("text")
                if page_text.strip():
                    text_parts.append(f"[Page {page_number}]\n{page_text}")
        return "\n\n".join(text_parts)

    @staticmethod
    def _extract_docx(file_path: Path) -> str:
        doc = DocxDocument(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        table_texts: list[str] = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    table_texts.append(" | ".join(cells))

        return "\n".join(paragraphs + table_texts)

    @staticmethod
    def _extract_txt(file_path: Path) -> str:
        for encoding in ("utf-8", "utf-16", "latin-1"):
            try:
                return file_path.read_text(encoding=encoding)
            except (UnicodeDecodeError, UnicodeError):
                continue
        raise ValueError("Unable to decode text file with supported encodings")

    @staticmethod
    def _extract_csv(file_path: Path) -> str:
        raw_bytes = file_path.read_bytes()
        text = raw_bytes.decode("utf-8", errors="ignore")
        reader = csv.reader(io.StringIO(text))
        rows = list(reader)
        if not rows:
            return ""

        header, data_rows = rows[0], rows[1:]
        lines = []
        for row in data_rows:
            paired = [f"{h.strip()}: {v.strip()}" for h, v in zip(header, row) if v.strip()]
            if paired:
                lines.append(", ".join(paired))
        return "\n".join(lines)
